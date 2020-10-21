import io

import requests
import pytesseract
from datetime import date

from PIL import Image
from docx import Document
from docx.shared import Cm

from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager



today = date.today()

open_file_read = open('pat_id_temp.txt', 'r')
pat_id_list = open_file_read.readlines()
open_file_read.close()

patient_details = pat_id_list[0]
patient_details = patient_details.split(' ')
pat_id = patient_details[0]
pat_doa = patient_details[1]

# --------------------v1.0- pytesseract and webdriver path need to be allocated--------------------------#
# pytesseract.pytesseract.tesseract_cmd = "C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
# pytesseract.pytesseract.tesseract_cmd = "C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe"
# driver = webdriver.Chrome(r'''C:\Users\AIIMS\Desktop\thidid\ths\drvers\chromedriver.exe''')
# driver = webdriver.Chrome(r'''C:\Users\DESKTOP\PycharmProjects\thidid\ths\drvers\chromedriver.exe''')
# -------------------------------------------------------------------------------------------------------#

# -----------------v2.0- webdriver path is extracted directly from the local directory-------------------#
# PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))
# DRIVER_BIN = os.path.join(PROJECT_ROOT, "chromedriver")
# driver = webdriver.Chrome(executable_path=DRIVER_BIN)
# -------------------------------------------------------------------------------------------------------#

# ----v3.0- webdriver path is automatically extracted from cache and is updated as per chrome version----#
driver = webdriver.Chrome(ChromeDriverManager().install())
# -------------------------------------------------------------------------------------------------------#

# URLs
main_url = 'http://212.83.179.41/~aiimshis/patient_record/opd_pat_rpt.php?pat_id=' + pat_id
page2 = main_url + r'''&page_no=2'''

# Creating the dictionary of investigations on patient's profile
# Page 1
thrfs = {}
driver.get(main_url)
hrfs1 = []
tastes1 = driver.find_elements_by_class_name('table-content')
buttons1 = driver.find_elements_by_class_name('button')
for i in reversed(buttons1):
    hrfs1.append(i.get_attribute('href'))
for i, j in zip(hrfs1, reversed(tastes1)):
    thrfs[j.text] = i

# Page 2
driver.get(page2)
tastes2 = driver.find_elements_by_class_name('table-content')
buttons2 = driver.find_elements_by_class_name('button')
hrfs2 = []
for i in reversed(buttons2):
    hrfs2.append(i.get_attribute('href'))
for i, j in zip(hrfs2, reversed(tastes2)):
    thrfs[j.text] = i


# Scrapping functions
def dater(a):
    return date(int(a[0:4]), int(a[5:7]), int(a[8:10]))


def findall(a):
    uil = a.find('Result')
    return a[uil + 6:len(a)]


def findernamer(raw, start, end):
    s = int(raw.find(start))
    e = int(raw.find(end))
    e1 = raw[s + len(start):e]
    e2 = e1.strip()
    return e2


def ocr_searcher(a, raw):
    return findnum(raw, a)


def findnum(a, b):
    a = a + " "
    i = a.find(b)
    ret = ''
    for x in a[i:]:

        if x.isdigit() == True or x == ".":
            ret = ret + x
        elif ret != '':
            return ret
    return ""


def find_val(a, b):
    return findnum(a, b)


# Required Dictionaries for storage of scrapped data
kft_dict = {}
lft_dict = {}
serum_e_dict = {}
cbc_dict = {}
hs_crp_dict = {}
esr_dict = {}
ret_dict = {}
ptinr_dict = {}
ldh_dict = {}
ca_dict = {}
phos_dict = {}
uric_dict = {}
general_dict = {}
e_thang_dict = {}


# Patient Details
class namer:
    def __init__(self):
        driver.get(main_url)
        n_page = driver.find_element_by_class_name('button').get_attribute('href')
        driver.get(n_page)
        r_text = driver.find_elements_by_class_name('content')[1].text
        self.pname = findernamer(r_text, r'''Patient's :''', r'''Father's / Spouse's :''')
        self.fname = findernamer(r_text, r'''Father's / Spouse's :''', r'''Reg. Date :''')
        self.rdate = findernamer(r_text, r'''Reg. Date :''', r'''Address :''')
        self.adr = findernamer(r_text, r'''Address :''', r'''Age :''')
        self.age = findernamer(r_text, r'''Age :''', r'''Gender :''')
        self.gender = findernamer(r_text, r'''Gender :''', r'''Contact Number :''')
        self.nom = findernamer(r_text, r'''Contact Number :''', r'''Valid From :''')

    def tabler(self):
        table = doc.tables[0]
        table.cell(0, 0).paragraphs[0].add_run(pat_id + ' designed by X').bold = True
        table.cell(1, 0).paragraphs[0].add_run(' ' + self.pname)
        table.cell(1, 1).paragraphs[0].add_run(' ' + self.fname)
        table.cell(1, 2).paragraphs[0].add_run(' ' + self.rdate)
        table.cell(2, 0).paragraphs[0].add_run(' ' + self.adr)
        table.cell(2, 1).paragraphs[0].add_run(' ' + self.age)
        table.cell(2, 2).paragraphs[0].add_run(' ' + self.gender)
        table.cell(3, 0).paragraphs[0].add_run(' ' + self.nom)
        table.cell(3, 1).paragraphs[0].add_run(' ' + pat_doa)
        table.cell(3, 2).paragraphs[0].add_run(' ' + str(today.day) + '-' + str(today.month) + '-' + str(today.year))


# Details of Complete Blood Count
class cbc:
    finder = ['Complete Haemogram', 'Emergency Haemogram']
    dik = cbc_dict
    table_i = 1

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        url = driver.find_element_by_tag_name('img').get_attribute('src')
        response = requests.get(url)
        img = Image.open(io.BytesIO(response.content))
        raw = pytesseract.image_to_string(img)
        self.hb = ocr_searcher("HGB", raw)
        self.tlc = ocr_searcher("WBC", raw)
        self.dlc = ocr_searcher("NEUT", raw) + '/' + ocr_searcher("LYMPH", raw) \
                   + '/' + ocr_searcher("MONO", raw) + '/' + ocr_searcher("EO", raw)
        self.plt = ocr_searcher("PLT", raw)
        self.hct = ocr_searcher("HCT", raw)
        self.mcv = ocr_searcher("MCV", raw)

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 1
        for i, j in cls.dik.items():
            col_cells = table.add_column(36000).cells
            for cell in table.columns[x].cells:
                cell.width = Cm(2.44)
            col_cells[0].paragraphs[0].add_run(
                str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]).bold = True
            col_cells[1].text = j.hb
            col_cells[2].text = j.tlc
            col_cells[3].text = j.dlc
            col_cells[4].text = j.plt
            col_cells[5].text = j.hct
            col_cells[6].text = j.mcv
            x = x + 1
        table.style = 'TableGrid'


# Details of Kidney Function Tests
class kft:
    finder = ['Kidney Function Test (KFT)']
    dik = kft_dict
    table_i = 2

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.urea = find_val(r_test, 'Urea')
        self.creat = find_val(r_test, 'Creatinine')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 1
        for i, j in cls.dik.items():
            col_cells = table.add_column(36000).cells
            for cell in table.columns[x].cells:
                cell.width = Cm(1.5)
            col_cells[0].paragraphs[0].add_run(
                str(today.day) + '-' + str(today.month) + '-' + str(today.year)).bold = True
            col_cells[1].text = j.urea
            col_cells[2].text = j.creat
            x = x + 1


# Details of Serum Electrolyte
class serum_e:
    finder = ['Chlorides (with Sodium Potassium)', 'Potassium/sodium']
    dik = serum_e_dict
    table_i = 3

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.na = ""
        self.k = ""
        self.cl = ""
        if r_test.find('Sodium') != -1:
            self.na = find_val(r_test, 'Sodium')
        if r_test.find('Potassium') != -1:
            self.k = find_val(r_test, 'Potassium')
        if r_test.find('Chloride') != -1:
            self.cl = find_val(r_test, 'Chloride')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 1
        for i, j in cls.dik.items():
            col_cells = table.add_column(36000).cells
            for cell in table.columns[x].cells:
                cell.width = Cm(1.23)
            col_cells[0].paragraphs[0].add_run(
                str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]).bold = True
            col_cells[1].text = j.na
            col_cells[2].text = j.k
            col_cells[3].text = j.cl
            x = x + 1


# Details of Liver Function Tests
class lft:
    finder = ['Liver Function Test (LFT)']
    dik = lft_dict
    table_i = 4

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.sgpt = find_val(r_test, 'SGPT')
        self.sgot = find_val(r_test, 'SGOT')
        self.bt = find_val(r_test, 'Total Bilirubin')
        self.bd = find_val(r_test, 'Direct Bilirubin')
        self.bi = find_val(r_test, 'Indirect Bilirubin')
        self.prt = find_val(r_test, 'Total Protein')
        self.alb = find_val(r_test, 'Albumin')
        self.glob = find_val(r_test, 'Globulin')
        self.alp = find_val(r_test, 'ALP')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 1
        for i, j in cls.dik.items():
            col_cells = table.add_column(36000).cells
            for cell in table.columns[x].cells:
                cell.width = Cm(3.13)
            col_cells[0].paragraphs[0].add_run(
                str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]).bold = True
            col_cells[1].text = j.sgpt + '/' + j.sgot + '/' + j.alp
            col_cells[2].text = j.bt + '/' + j.bd + '/' + j.bi
            col_cells[3].text = j.prt + '/' + j.alb + '/' + j.glob
            x = x + 1


# Details of Prothrombin Time and International Normalised Ratio
class ptinr:
    finder = ['Prothrombin Time (PT) & IN']
    dik = ptinr_dict
    table_i = 5

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.pt = find_val(r_test, 'Patient Time')
        self.inr = find_val(r_test, 'INR ')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.pt + '/' + j.inr
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.pt + '/' + j.inr
            x = x + 1


# Details of Erythrocyte Sedimentation Rate
class esr:
    finder = ['Erythrocyte Sedimentation Rate']
    dik = esr_dict
    table_i = 6

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.esr_ = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.esr_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.esr_
            x = x + 1


# Details of Hs- C Reactive Protein
class hs_crp:
    finder = ['Hs CRP']
    dik = hs_crp_dict
    table_i = 7

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.crp = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.crp
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.crp
            x = x + 1


# Details of Calcium levels
class ca:
    finder = ['Calcium']
    dik = ca_dict
    table_i = 8

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.ca_ = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.ca_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.ca_
            x = x + 1


# Details of Phosphate Levels
class phos:
    finder = ['Phosphorus']
    dik = phos_dict
    table_i = 9

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.phos_ = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.phos_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.phos_
            x = x + 1


# Details of Reticulocyte count
class ret:
    finder = ['Reticulocyte Count']
    dik = ret_dict
    table_i = 10

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.count_ = find_val(r_test, 'count is')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.count_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.count_
            x = x + 1


# Details of LDH
class ldh:
    finder = ['LDH']
    dik = ldh_dict
    table_i = 11

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.ldh_ = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.ldh_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.ldh_
            x = x + 1


# Details of Serum Uric Acid
class uric:
    finder = ['Serum Uric Acid']
    dik = uric_dict
    table_i = 12

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.uric_ = find_val(r_test, 'Result')

    @staticmethod
    def tabler():
        table = doc.tables[1].rows[cls.table_i - 1].cells[1].tables[0]
        x = 0
        for i, j in cls.dik.items():
            if x == 0:
                col_cells = table.rows[0].cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.uric_
            else:
                col_cells = table.add_row().cells
                col_cells[0].text = str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]
                col_cells[1].text = j.uric_
            x = x + 1


# Other general details
class general:
    finder = ['HIV1/2 Rapid Test', 'HCV Antibody Rapid', 'HBsAg Rapid',
              'HbA1c (Glycated Haemoglobin)',
              'Urine Microscopy',
              'Iron', 'TIBC', 'Ferritin',
              'Thyroid Stimulating Hormone (TSH)',
              'Activated Partial Thromboplastin Time (APTT)',
              'Procalcitonin', 'Lipid Profile (Ch+TG+LDL+HDL+calc)']
    dik = general_dict
    table_i = 13

    def __init__(self):
        self.date = dater(driver.find_element_by_class_name('table-header').text)
        self.tname = driver.find_element_by_class_name('table-content').text.strip('Test Name')
        r_test = driver.find_elements_by_class_name('table-content')[1].text
        self.all_ = r_test.strip('Result').strip()

    @staticmethod
    def tabler():
        x = -1
        for i, j in cls.dik.items():
            r_cells = doc.tables[1].rows[cls.table_i + x].cells
            r_cells[0].paragraphs[0].add_run(j.tname).bold = True
            r_cells[1].paragraphs[0].add_run(
                str(j.date.day) + '/' + str(j.date.month) + '/' + str(j.date.year)[2:4]).bold = True
            r_cells[2].paragraphs[0].add_run(j.all_)
            x = x + 1


master_classes = [kft, cbc, lft, serum_e, hs_crp, esr, ret, ptinr, ldh, ca, phos, general, uric]
namer1 = namer()
for cls in master_classes:
    for i, j in thrfs.items():
        for x in cls.finder:
            if x in i:
                driver.get(j)
                cls.dik[i] = cls()
                # if dater(driver.find_element_by_class_name('table-header').text) > pat_doa:
                # cls.dik[i] = cls()

doc = Document('model.docx')
namer1.tabler()
for cls in master_classes:
    cls.tabler()
driver.close()
name = namer1.pname
doc.save(namer1.pname + '.docx')
